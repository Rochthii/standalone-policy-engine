import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import win32com.client

def set_cell_background(cell, fill_color):
    """Đặt màu nền cho ô trong bảng (hex, vd: '1E3A5F')"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Đặt padding cho ô trong bảng"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_callout(doc, text, title="LƯU Ý CỐT LÕI (FIRST PRINCIPLES)", bg_color="F0F4F8", border_color="1E3A5F"):
    """Tạo hộp Callout / Note chuyên nghiệp"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border highlight
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    if title:
        run_title = p.add_run(f"📌 {title}\n")
        run_title.bold = True
        run_title.font.name = "Arial"
        run_title.font.size = Pt(10.5)
        run_title.font.color.rgb = RGBColor(30, 58, 95)
    
    run_text = p.add_run(text)
    run_text.font.name = "Arial"
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(40, 50, 60)
    
    # Empty spacing after table
    sp_p = doc.add_paragraph()
    sp_p.paragraph_format.space_before = Pt(0)
    sp_p.paragraph_format.space_after = Pt(4)

def format_heading(p, text, level=1):
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.bold = True
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(14, 43, 82) # Deep Blue
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(12.5)
        run.font.color.rgb = RGBColor(13, 79, 60) # Deep Emerald
    elif level == 3:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(50, 60, 75)

def format_bullet(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    r1 = p.add_run(bold_prefix)
    r1.font.name = "Arial"
    r1.font.size = Pt(10)
    r1.bold = True
    r1.font.color.rgb = RGBColor(20, 30, 40)
    
    r2 = p.add_run(text)
    r2.font.name = "Arial"
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(50, 60, 70)

def main():
    doc = Document()
    
    # Page Margins: 1 inch (72 pt) all around
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Header / Footer
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Bản Quyền Lộ Trình SE & ERP (SAP / Odoo) — Standalone PDP Engine Project")
        f_run.font.name = "Arial"
        f_run.font.size = Pt(8)
        f_run.font.color.rgb = RGBColor(140, 150, 160)

    # ----------------------------------------------------
    # COVER / HEADER BANNER
    # ----------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    t_run = title_p.add_run("LỘ TRÌNH CĂN TÍNH:\nKỸ THUẬT PHẦN MỀM & KIẾN TRÚC ERP (SAP & ODOO)\nTRONG KỶ NGUYÊN AI VIBE CODING")
    t_run.font.name = "Arial"
    t_run.font.size = Pt(18)
    t_run.bold = True
    t_run.font.color.rgb = RGBColor(14, 43, 82)
    
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(2)
    sub_p.paragraph_format.space_after = Pt(14)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    s_run = sub_p.add_run("Bản Thiết Kế Toàn Diện Từ Năm 3 Đại Học Đến Vị Trí Enterprise Solutions Architect")
    s_run.font.name = "Arial"
    s_run.font.size = Pt(11)
    s_run.italic = True
    s_run.font.color.rgb = RGBColor(80, 95, 110)

    add_callout(
        doc,
        "Lộ trình này được xây dựng trên 3 trục Căn Tính Cốt Lõi (First Principles):\n"
        "1. CĂN TÍNH KINH TẾ (ERP Domain): Hiểu thấu đáo cách tiền, tài sản, kho bãi và báo cáo tài chính vận hành trong thế giới thực.\n"
        "2. CĂN TÍNH HỆ THỐNG (SE Rigor): Kiến trúc vi dịch vụ phân tán, gRPC, OData, ACID, In-Memory Engine, Copy-On-Write lock-free.\n"
        "3. CĂN TÍNH THỜI ĐẠI (AI Agentic Orchestration): Bạn là Kiến trúc sư trưởng đưa ra đặc tả kỹ thuật; AI là trợ lý lập trình thực thi.",
        title="TRIẾT LÝ PHÁT TRIỂN NĂNG LỰC CỐT LÕI"
    )

    # ----------------------------------------------------
    # CHƯƠNG 1: TỔNG QUAN & BẢN ĐỒ CHIẾN LƯỢC
    # ----------------------------------------------------
    format_heading(doc.add_paragraph(), "1. Tổng Quan 4 Giai Đoạn & So Sánh Chiến Lược SAP vs Odoo", level=1)
    
    # Table 1: So sánh SAP vs Odoo
    tbl_cmp = doc.add_table(rows=5, cols=3)
    tbl_cmp.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_cmp.autofit = False
    
    col_widths = [Inches(1.8), Inches(2.4), Inches(2.4)]
    headers = ["Tiêu Chí Đánh Giá", "Hệ Sinh Thái Odoo ERP", "Hệ Sinh Thái SAP (S/4HANA / BTP)"]
    
    # Header Row
    hdr_row = tbl_cmp.rows[0]
    for i, title in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = col_widths[i]
        set_cell_background(cell, "1E3A5F")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    rows_data = [
        ("Phân khúc thị trường", "Doanh nghiệp vừa & nhỏ (SME), Tech startup", "Tập đoàn tỷ USD, Đa quốc gia (Vinamilk, Big 4)"),
        ("Công nghệ & Mã nguồn", "100% Open-Source (Python, PostgreSQL, OWL)", "Enterprise Cloud (ABAP on Cloud, SAP BTP, OData)"),
        ("Mục tiêu học tập", "Khai phá ruột gan code, hiểu cấu trúc dữ liệu", "Học chuẩn quy trình doanh nghiệp toàn cầu"),
        ("Vị trí & Thu nhập", "Odoo Developer / Lead ($800 - $2,500)", "SAP Consultant / Architect ($1,500 - $5,000+)")
    ]
    
    for r_idx, row_data in enumerate(rows_data):
        row = tbl_cmp.rows[r_idx + 1]
        bg = "F7FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 60, 60, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Arial"
            r.font.size = Pt(9)
            if c_idx == 0:
                r.bold = True
            r.font.color.rgb = RGBColor(30, 40, 50)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ----------------------------------------------------
    # CHƯƠNG 2: GIAI ĐOẠN 1 - ODOO & LẬP TRÌNH PHÂN TÁN
    # ----------------------------------------------------
    format_heading(doc.add_paragraph(), "2. Giai Đoạn 1: Làm Chủ Ruột Gan ERP Với Odoo & Lập Trình Phân Tán", level=1)
    
    p_g1 = doc.add_paragraph()
    p_g1.paragraph_format.space_after = Pt(4)
    r = p_g1.add_run("Thời gian: Năm 3 (Tháng 1 — Tháng 3) | Mục tiêu: Hiểu thấu 100% cách ERP tổ chức dữ liệu & dòng tiền.")
    r.font.name = "Arial"; r.font.size = Pt(10); r.italic = True

    format_heading(doc.add_paragraph(), "A. Nội Dung Cần Học (Kiến Thức & Nghiệp Vụ Chuẩn):", level=2)
    format_bullet(doc, "4 Luồng Nghiệp Vụ Cốt Lõi: ", "P2P (Procure-to-Pay: PR → PO → GRN → Vendor Bill → Payment), O2C (Order-to-Cash: Báo giá → SO → Xuất kho → Invoice → Thu tiền), Kế toán kép (Double-entry Debit/Credit), Kho vận (Double-entry Inventory Moves, FIFO/AVCO).")
    format_bullet(doc, "Kỹ Thuật Odoo ORM: ", "Kế thừa Model (_inherit), quan hệ Many2one/One2many, Decorators (@api.depends, @api.constrains), Phân quyền Record Rules (ir.rule).")
    format_bullet(doc, "Database PostgreSQL: ", "Cơ chế Transaction ACID, Row-level Locks, MVCC và cách Odoo tối ưu truy vấn.")

    format_heading(doc.add_paragraph(), "B. Nhiệm Vụ Cần Làm (Thực Chiến Checklist):", level=2)
    format_bullet(doc, "Task 1.1: ", "Viết docker-compose.yml chạy cụm Odoo 17 Community + PostgreSQL 16 cục bộ.")
    format_bullet(doc, "Task 1.2: ", "Thao tác trọn vẹn 1 chu trình mua hàng P2P và 1 chu trình bán hàng O2C trên giao diện Odoo thật.")
    format_bullet(doc, "Task 1.3: ", "Viết Custom Odoo Module bằng Python: Khi bấm 'Approve PO', gửi request gRPC/REST sang Standalone Go PDP Engine để đánh giá quyền ABAC.")

    format_heading(doc.add_paragraph(), "C. Sản Phẩm Đầu Ra & Điểm Nhấn CV:", level=2)
    format_bullet(doc, "Sản phẩm: ", "Repository GitHub chứa module Odoo tích hợp Standalone Go PDP hoàn chỉnh.")
    format_bullet(doc, "Bullet Point CV: ", "'Engineered custom Odoo 17 enterprise modules with distributed gRPC authorization hooks, offloading dynamic ABAC decision logic to an in-memory Go Policy Decision Point with sub-millisecond evaluation latency.'")

    # ----------------------------------------------------
    # CHƯƠNG 3: GIAI ĐOẠN 2 - CHUẨN MỰC SAP ENTERPRISE
    # ----------------------------------------------------
    format_heading(doc.add_paragraph(), "3. Giai Đoạn 2: Tiếp Cận Chuẩn Mực Doanh Nghiệp Tỷ USD Với SAP", level=1)
    
    p_g2 = doc.add_paragraph()
    p_g2.paragraph_format.space_after = Pt(4)
    r = p_g2.add_run("Thời gian: Hè Năm 3 (Tháng 4 — Tháng 6) | Mục tiêu: Nắm vững chuẩn mực quy trình SAP S/4HANA, SAP BTP & Kiến trúc Two-Tier ERP.")
    r.font.name = "Arial"; r.font.size = Pt(10); r.italic = True

    format_heading(doc.add_paragraph(), "A. Nội Dung Cần Học:", level=2)
    format_bullet(doc, "Cấu Trúc Doanh Nghiệp SAP: ", "Client → Company Code (BUKRS - Pháp nhân) → Plant (WERKS - Nhà máy) → Purchasing/Sales Org → Cost Center (KOSTL).")
    format_bullet(doc, "Công Nghệ SAP Clean Core: ", "Nguyên tắc mở rộng Side-by-Side trên SAP BTP, giao thức chuẩn SAP OData v2/v4 Services (JSON/REST APIs).")
    format_bullet(doc, "Bảo Mật & Phân Quyền SAP GRC: ", "Authorization Objects (M_BEST_EKO, F_BKPF_BUK) và Ma trận Phân tách trách nhiệm (Separation of Duties - SoD).")

    format_heading(doc.add_paragraph(), "B. Nhiệm Vụ Cần Làm:", level=2)
    format_bullet(doc, "Task 2.1: ", "Học và hoàn thành chứng chỉ miễn phí trên learning.sap.com / openSAP (Discovering SAP S/4HANA, SAP BTP).")
    format_bullet(doc, "Task 2.2: ", "Dùng AI viết Mock SAP S/4HANA OData Server bằng Go (EntitySet /A_PurchaseOrder, /A_JournalEntry).")
    format_bullet(doc, "Task 2.3 (Two-Tier ERP Bridge): ", "Đồng bộ giao dịch PO/Invoice từ Odoo (Chi nhánh) sang Mock SAP S/4HANA (Trụ sở tập đoàn), bảo vệ an toàn bởi Go PDP Engine.")

    format_heading(doc.add_paragraph(), "C. Sản Phẩm Đầu Ra & Điểm Nhấn CV:", level=2)
    format_bullet(doc, "Sản phẩm: ", "1-2 Chứng chỉ chính thức từ SAP Official + Mô hình Two-Tier ERP tích hợp.")
    format_bullet(doc, "Bullet Point CV: ", "'Architected an event-driven Two-Tier ERP integration bridge syncing Branch transactions (Odoo) to Corporate HQ (SAP S/4HANA OData v4 APIs) guarded by a high-throughput Go ABAC Policy Engine.'")

    # ----------------------------------------------------
    # CHƯƠNG 4: GIAI ĐOẠN 3 - ĐỒ ÁN TỐT NGHIỆP & SĂN THỰC TẬP
    # ----------------------------------------------------
    format_heading(doc.add_paragraph(), "4. Giai Đoạn 3: Hoàn Thiện Đồ Án Tốt Nghiệp Xuất Sắc & Săn Thực Tập", level=1)
    
    p_g3 = doc.add_paragraph()
    p_g3.paragraph_format.space_after = Pt(4)
    r = p_g3.add_run("Thời gian: Năm 4 (Tháng 7 — Tháng 12) | Mục tiêu: Đạt điểm Đồ án 9.5 - 10 & Nhận Offer Thực tập sinh/Fresher có lương.")
    r.font.name = "Arial"; r.font.size = Pt(10); r.italic = True

    format_heading(doc.add_paragraph(), "A. Kịch Bản Live Demo 5 Phút Trước Hội Đồng:", level=2)
    format_bullet(doc, "Phút 1: ", "Trình bày sơ đồ kiến trúc tổng thể (Envoy PEP ↔ PDP Go Server :50051 ↔ PostgreSQL).")
    format_bullet(doc, "Phút 2: ", "Demo trực tiếp trên giao diện Odoo: Duyệt PO đúng hạn mức (ALLOW) vs Người tạo tự duyệt (DENY do SoD).")
    format_bullet(doc, "Phút 3: ", "Demo tính năng ExplainDecision gRPC trả về chính xác Policy ID phục vụ kiểm toán.")
    format_bullet(doc, "Phút 4: ", "Demo Hot-Reload < 300ms qua Redis Pub/Sub và cơ chế Spill-to-Disk khi Database PostgreSQL bị ngắt kết nối.")
    format_bullet(doc, "Phút 5: ", "Chạy Benchmark trực tiếp: Chứng minh thông lượng > 1.000.000 RPS và độ trễ 0.9 µs trong RAM.")

    format_heading(doc.add_paragraph(), "B. Danh Sách Doanh Nghiệp Mục Tiêu Nộp CV:", level=2)
    format_bullet(doc, "Doanh nghiệp Triển khai ERP lớn: ", "FPT Software (FJP/FHN/FHM), Viettel Solutions, CMC Global, SmartOSC, BAP, A1 Consulting.")
    format_bullet(doc, "Tập đoàn Tư vấn Quốc tế (Big 4): ", "Deloitte, PwC, EY, KPMG (Technology Consulting / Technology Risk).")

    # ----------------------------------------------------
    # CHƯƠNG 5: GIAI ĐOẠN 4 - ENTERPRISE SOLUTIONS ARCHITECT
    # ----------------------------------------------------
    format_heading(doc.add_paragraph(), "5. Giai Đoạn 4: Nâng Tầm Thành Enterprise Solutions Architect", level=1)
    
    format_heading(doc.add_paragraph(), "Lộ Trình Thăng Tiến 3 Năm:", level=2)
    format_bullet(doc, "Năm 1: Technical Consultant (Thu nhập $800 - $1,500/tháng) ", "Tham gia dự án triển khai SAP/Odoo thực tế, dùng AI tăng năng suất viết module và tích hợp API gấp 3-5 lần.")
    format_bullet(doc, "Năm 2-3: Senior Integration Lead (Thu nhập $1,800 - $3,000/tháng) ", "Thiết kế kiến trúc tích hợp hệ thống lớn, tối ưu hóa Database & Bảo mật, lấy chứng chỉ SAP Certified Development & AWS Solutions Architect.")
    format_bullet(doc, "Năm 3+: Principal Solutions Architect (Thu nhập $3,500 - $6,000+/tháng) ", "Định hình kiến trúc chuyển đổi số tổng thể cho tập đoàn, quản trị rủi ro Zero Trust.")

    # ----------------------------------------------------
    # CHƯƠNG 6: BẢNG SỐ LIỆU THỰC NGHIỆM ĐỒ ÁN
    # ----------------------------------------------------
    format_heading(doc.add_paragraph(), "6. Bảng Đo Lường Hiệu Năng Thực Nghiệm Standalone PDP Engine", level=1)
    
    tbl_bench = doc.add_table(rows=4, cols=5)
    tbl_bench.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_bench.autofit = False
    
    b_col_widths = [Inches(2.2), Inches(1.1), Inches(1.0), Inches(1.0), Inches(1.3)]
    b_headers = ["Kịch Bản Kiểm Thử", "Độ Trễ / Op", "RAM B/op", "Alloc/op", "Throughput Ước Tính"]
    
    hdr_row = tbl_bench.rows[0]
    for i, title in enumerate(b_headers):
        cell = hdr_row.cells[i]
        cell.width = b_col_widths[i]
        set_cell_background(cell, "0D4F3C") # Deep Emerald
        set_cell_margins(cell, 80, 80, 80, 80)
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.font.name = "Arial"; r.font.size = Pt(9); r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255)
        
    b_data = [
        ("PO Approval ABAC (Đơn luồng)", "5.83 µs/op", "1,454 B", "33 allocs", "~171,000 req/s/core"),
        ("Đa luồng đồng thời (20 Cores)", "946.5 ns/op", "748 B", "11 allocs", "> 1,050,000 req/s"),
        ("Tra cứu Chỉ mục Trie thuần túy", "449.0 ns/op", "325 B", "10 allocs", "> 2,200,000 req/s")
    ]
    
    for r_idx, row_data in enumerate(b_data):
        row = tbl_bench.rows[r_idx + 1]
        bg = "F0FFF4" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = b_col_widths[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 60, 60, 80, 80)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Arial"; r.font.size = Pt(8.5)
            if c_idx == 0 or c_idx == 1 or c_idx == 4:
                r.bold = True
            r.font.color.rgb = RGBColor(20, 35, 25)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Save DOCX
    docx_path = r"e:\Projects\Project_TN\standalone-policy-engine\docs\career-roadmap\LO_TRINH_SE_ERP_AI_VIBE_CODING.docx"
    doc.save(docx_path)
    print(f"DOCX created successfully at: {docx_path}")

    # Export to PDF via Word COM
    pdf_path = r"e:\Projects\Project_TN\standalone-policy-engine\docs\career-roadmap\LO_TRINH_SE_ERP_AI_VIBE_CODING.pdf"
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc_com = word.Documents.Open(os.path.abspath(docx_path))
        # 17 = wdExportFormatPDF
        doc_com.ExportAsFixedFormat(os.path.abspath(pdf_path), 17)
        doc_com.Close(False)
        word.Quit()
        print(f"PDF exported successfully at: {pdf_path}")
    except Exception as e:
        print(f"Error exporting PDF via Word COM: {e}")

if __name__ == "__main__":
    main()
